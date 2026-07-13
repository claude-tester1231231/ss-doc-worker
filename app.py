# ss-doc-worker — document extract/rebuild engine (PyMuPDF + python-docx).
#
# This service deliberately contains NO translation logic, NO prompts, NO API
# keys and NO business rules. It does exactly two things per format:
#   extract: document in  -> ordered text blocks (id + text) out
#   build:   document + {id: replacement-text} in -> rebuilt document out
# The caller owns everything else. That separation is the point: the whole
# PyMuPDF-touching surface lives in this small AGPL-licensed service.
#
# Auth: shared bearer token from WORKER_TOKEN env (config, not code).
import io
import json
import os
import re

import fitz  # PyMuPDF (AGPL)
import docx  # python-docx (MIT)
from docx.oxml.ns import qn
from flask import Flask, abort, jsonify, request, send_file

app = Flask(__name__)
MAX_BYTES = 25 * 1024 * 1024
MIN_FONT = 8.0

MARK = {(True, False): 'B', (False, True): 'I', (True, True): 'BI'}
M_RE = re.compile(r'\[\[(/?)(BI|B|I)\]\]')


def _auth():
    want = os.environ.get('WORKER_TOKEN', '')
    got = (request.headers.get('Authorization') or '').replace('Bearer ', '', 1)
    if not want or got != want:
        abort(401)


def _file_bytes(field='file'):
    f = request.files.get(field)
    if f is None:
        abort(400, 'missing file')
    data = f.read()
    if not data or len(data) > MAX_BYTES:
        abort(413 if data else 400)
    return data


# ── PDF lane ────────────────────────────────────────────────────────────────

BULLET_RE = re.compile(r'^\s*(?:[•▪◦–\-\*]|\d{1,2}[.)])\s')
# a bullet glyph alone on its line (PyMuPDF often splits '• text' into two
# lines) — folded into the NEXT line so list rows keep their original height
BULLET_ONLY_RE = re.compile(r'^\s*[•▪◦–\-\*]\s*$')


def _font_class(name):
    n = (name or '').lower()
    if 'mono' in n or 'courier' in n or 'consol' in n:
        return 'mono'
    if 'times' in n or 'serif' in n or 'georgia' in n or 'garamond' in n or 'book' in n:
        return 'serif'
    return 'sans'


def pdf_collect(doc):
    """The document's OWN text blocks -> paragraph-merged jobs. Line-aware:
    lines join with newlines (bullet lists survive translation as lines — live
    12/7: space-joins turned lists into one clump). Dominant font family rides
    along so the rebuild can match sans/serif/mono. Bullet blocks never merge."""
    jobs = []
    for pno, page in enumerate(doc):
        raw = []
        for b in page.get_text('dict')['blocks']:
            if b.get('type') != 0:
                continue
            # collect (y, text) per line-object; PyMuPDF emits side-by-side text
            # (form label + value) as SEPARATE line-objects at the SAME y —
            # joining those with \n forces 2x height into a 1-line box (live 12/7
            # tight_boxes bug). Group by vertical position: same row -> space,
            # real vertical gap -> newline.
            fonts = []
            rows = []  # [ (y0, height, text) ]
            for l in b.get('lines', []):
                spans = l.get('spans', [])
                lt = ' '.join(s['text'] for s in spans).strip()
                fonts.extend(s.get('font', '') for s in spans)
                if not lt:
                    continue
                lb = l.get('bbox', b['bbox'])
                ly, lh = lb[1], max(1.0, lb[3] - lb[1])
                if rows and abs(ly - rows[-1][0]) < 0.6 * rows[-1][1]:
                    rows[-1] = (rows[-1][0], rows[-1][1], rows[-1][2] + ' ' + lt)
                else:
                    rows.append((ly, lh, lt))
            if not rows:
                continue
            rows.sort(key=lambda r: r[0])
            lines = [r[2] for r in rows]
            folded = []
            for lt in lines:
                if folded and BULLET_ONLY_RE.match(folded[-1]):
                    folded[-1] = folded[-1].strip() + ' ' + lt
                else:
                    folded.append(lt)
            lines = folded
            spans = [s for l in b.get('lines', []) for s in l.get('spans', [])]
            dom_font = max(set(fonts), key=fonts.count) if fonts else ''
            raw.append({
                'page': pno, 'rect': fitz.Rect(b['bbox']), 'text': '\n'.join(lines),
                'size': max(s['size'] for s in spans),
                'color': max(set(s['color'] for s in spans),
                             key=[s['color'] for s in spans].count),
                'bold': any('bold' in s.get('font', '').lower() for s in spans),
                'font': _font_class(dom_font),
            })
        raw.sort(key=lambda j: (j['rect'].y0, j['rect'].x0))
        merged = []
        for j in raw:
            m = merged[-1] if merged else None
            is_bullet = bool(BULLET_RE.match(j['text']))
            if m and m['page'] == j['page'] and not is_bullet:
                vgap = j['rect'].y0 - m['rect'].y1
                xover = min(m['rect'].x1, j['rect'].x1) - max(m['rect'].x0, j['rect'].x0)
                same_style = (m['bold'] == j['bold'] and abs(m['size'] - j['size']) < 0.6
                              and m['font'] == j['font'])
                if vgap < 0.9 * j['size'] and xover > 0 and same_style:
                    m['rect'] |= j['rect']
                    m['text'] += '\n' + j['text']
                    continue
            merged.append(dict(j))
        jobs.extend(merged)
    for i, j in enumerate(jobs):
        j['id'] = 'b%d' % i
    return jobs


@app.post('/v1/pdf/extract')
def pdf_extract():
    _auth()
    doc = fitz.open(stream=_file_bytes(), filetype='pdf')
    jobs = pdf_collect(doc)
    # 'page' (1-based) rides along so callers can pipeline per page (lazy view)
    return jsonify({
        'engine': 'pdf', 'pages': len(doc),
        'blocks': [{'id': j['id'], 'text': j['text'], 'page': j['page'] + 1} for j in jobs],
    })


def _actualtext_wrap(doc, page, pre_xrefs, text):
    """ToUnicode-armor (13/7, live-batteri-fund): MuPDF's htmlbox mapper
    kontekst-substituerede glyffer FORKERT i ToUnicode — thai stablede
    tonemærker blev '2'/'1' (นี้->นี2), arabiske lam-alef-ligaturer blev
    garbage (hvert «al-»-ord), latin fi/ff blev ligatur-kodepunkter. Glyfferne
    TEGNES korrekt; kun kopiér/markér/oplæs var ødelagt. Fix: hver bloks
    nyindsatte content-stream pakkes i marked-content med /ActualText =
    den logiske oversættelses-streng — spec-tro ekstraktorer (MuPDF, pdf.js,
    Acrobat) returnerer strengen verbatim, uafhængigt af glyf-mapping.
    8 trailing spaces absorberer rest-glyffer (glyfantal > tegnantal ved
    variant-marks); overskud kollapser som whitespace i ekstraktion."""
    at = '<FEFF' + (text + ' ' * 8).encode('utf-16-be').hex().upper() + '>'
    for x in page.get_contents():
        if x in pre_xrefs:
            continue
        raw = doc.xref_stream(x)
        doc.update_stream(
            x,
            b'/Span <</ActualText ' + at.encode('ascii') + b'>> BDC\n'
            + raw + b'\nEMC\n')


def _ocr_jobs(doc, raw):
    """OCR-lanen (13/7): scannede sider har ingen tekst-objekter — klienten
    sender Vision-paragraffer som normaliserede bokse [{id,page(1-based),
    x,y,w,h}] (0-1 af sidens rekt). Vi bygger jobs i pdf_collects form, så
    resten af build-loopet (headroom, htmlbox, ActualText, rapport) er ETT
    fælles spor. Redaktioner springes over af kalderen: kildeteksten er
    PIXELS — den dækkes i stedet med en hvid plade pr. blok (whiteout)."""
    try:
        blocks = json.loads(raw)
    except ValueError:
        abort(400, 'bad ocr_blocks')
    if not isinstance(blocks, list) or not blocks or len(blocks) > 800:
        abort(400, 'bad ocr_blocks')
    jobs = []
    for b in blocks:
        if not isinstance(b, dict) or not isinstance(b.get('id'), str):
            abort(400, 'bad ocr_blocks')
        try:
            pno = int(b['page']) - 1
            x, y, w, h = float(b['x']), float(b['y']), float(b['w']), float(b['h'])
        except (KeyError, TypeError, ValueError):
            abort(400, 'bad ocr_blocks')
        if pno < 0 or pno >= len(doc):
            abort(400, 'bad ocr_blocks')
        if not (0 <= x <= 1 and 0 <= y <= 1 and 0 < w <= 1 and 0 < h <= 1):
            abort(400, 'bad ocr_blocks')
        pr = doc[pno].rect
        rect = fitz.Rect(pr.x0 + x * pr.width, pr.y0 + y * pr.height,
                         pr.x0 + min(x + w, 1.0) * pr.width,
                         pr.y0 + min(y + h, 1.0) * pr.height)
        # skriftstørrelse fra boks-højden: OCR kender ingen fonte. En Vision-
        # paragraf kan være flerlinjet — antag linjehøjde ~1.35 og afled fra
        # ordboksenes typiske højde er ikke tilgængelig her, så clamp bredt;
        # htmlbox skalerer selv NED hvis oversættelsen er længere.
        lines = max(1, round(rect.height / max(10.0, rect.width * 0.06)))
        size = max(7.0, min(22.0, (rect.height / lines) * 0.72))
        jobs.append({'id': b['id'], 'page': pno, 'rect': rect, 'text': '',
                     'size': size, 'color': 0, 'bold': False, 'font': 'sans',
                     '_ocr': True})
    return jobs


@app.post('/v1/pdf/build')
def pdf_build():
    _auth()
    tr = json.loads(request.form.get('translations') or '{}')
    if not isinstance(tr, dict) or not tr:
        abort(400, 'missing translations')
    doc = fitz.open(stream=_file_bytes(), filetype='pdf')
    ocr_raw = request.form.get('ocr_blocks')
    ocr_mode = ocr_raw is not None and str(ocr_raw).strip() != ''
    jobs = _ocr_jobs(doc, ocr_raw) if ocr_mode else pdf_collect(doc)
    # optional page scope (1-based): build ONLY that page and return a 1-page
    # PDF — the lazy per-page pipeline. Ids stay the global b<N> ids.
    only = request.form.get('page')
    only_idx = None
    if only is not None and str(only).strip() != '':
        try:
            only_idx = int(only) - 1
        except ValueError:
            abort(400, 'bad page')
        if only_idx < 0 or only_idx >= len(doc):
            abort(400, 'bad page')
        jobs = [j for j in jobs if j['page'] == only_idx]
    # hard 1:1 contract: every block IN SCOPE must have a non-empty replacement
    missing = [j['id'] for j in jobs if not isinstance(tr.get(j['id']), str) or not tr[j['id']].strip()]
    if missing:
        abort(422, 'missing ids: ' + ','.join(missing[:20]))
    report = {'blocks': len(jobs), 'placed': 0, 'shrunk': 0, 'failed': 0}
    for page in doc:
        if only_idx is not None and page.number != only_idx:
            continue
        pj = [j for j in jobs if j['page'] == page.number]
        if not pj:
            continue
        # HEADROOM (hardening 12/7): a longer translation must flow DOWN into the
        # empty space below its block instead of overflowing its original height.
        # For each block, find the nearest block below it that horizontally
        # overlaps; the insert box may grow to just above that (or the bottom
        # margin). Short text stays put (htmlbox is top-aligned); long text uses
        # the whitespace the source left. Column-aware via the x-overlap test, so
        # a sidebar block never grows into a main-column block.
        bottom_margin = page.rect.height - 30
        for j in pj:
            limit = bottom_margin
            for k in pj:
                if k is j or k['rect'].y0 < j['rect'].y1 - 1:
                    continue
                xov = min(j['rect'].x1, k['rect'].x1) - max(j['rect'].x0, k['rect'].x0)
                if xov > 1 and k['rect'].y0 < limit:
                    limit = k['rect'].y0
            j['_grow_bottom'] = max(j['rect'].y1, limit - 2)
        if ocr_mode:
            # scannet side: kildeteksten er pixels — redaktion kan intet
            # fjerne. Hvid plade pr. blok (let padding) dækker originalen,
            # oversættelsen tegnes ovenpå. Standard OCR-replace-udseende.
            for j in pj:
                r = j['rect']
                pad = fitz.Rect(r.x0 - 1.5, r.y0 - 1.5, r.x1 + 1.5, r.y1 + 1.5)
                page.draw_rect(pad, color=None, fill=(1, 1, 1))
        else:
            for j in pj:
                page.add_redact_annot(j['rect'])
            page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)
        for j in pj:
            text = tr[j['id']]
            # insert_htmlbox: full Unicode (Cyrillic/Arabic/CJK, shaping, RTL) —
            # Base14 'helv' turned every non-Latin char into '?' (live 12/7).
            # The replacement text is DATA: escape it, never interpret as HTML.
            esc = (text.replace('&', '&amp;').replace('<', '&lt;')
                       .replace('>', '&gt;').replace('\n', '<br>'))
            color = '#%06x' % j['color']
            size = max(MIN_FONT, j['size'])
            fam = {'serif': 'Times, serif', 'mono': 'Courier, monospace'}.get(
                j.get('font', 'sans'), 'Helvetica, Arial, sans-serif')
            # word-break: long compounds (German "Kraftfahrzeug…") must wrap
            # instead of overflowing the box horizontally.
            base = ('font-size:%.1fpx;color:%s;font-family:%s;'
                    'overflow-wrap:anywhere;word-break:break-word' % (size, color, fam))
            if j['bold']:
                html = '<b style="%s">%s</b>' % (base, esc)
            else:
                html = '<span style="%s">%s</span>' % (base, esc)
            if j.get('_ocr'):
                # OCR: boksen ER den visuelle sandhed på en scannet side —
                # x-vækst ville tegne hen over nabospaltens PIXELS (ingen
                # whiteout dér). Kun nedad-vækst (grow_bottom er kolonne-
                # bevidst); htmlbox skalerer ned hvis det stadig er trangt.
                rect = fitz.Rect(j['rect'].x0, j['rect'].y0,
                                 j['rect'].x1 + 2,
                                 max(j['rect'].y1, j['_grow_bottom']))
            else:
                rect = fitz.Rect(j['rect'].x0, j['rect'].y0 - 0.15 * j['size'],
                                 j['rect'].x1 + 0.35 * j['rect'].width,
                                 max(j['rect'].y1 + 0.45 * j['size'], j['_grow_bottom']))
            rect.x1 = min(rect.x1, page.rect.width - 8)
            # NEVER-VANISH (Morten 12/7 "text just disappears"): scale_low is a
            # low floor so insert_htmlbox picks the LARGEST scale in [low,1] that
            # fits. Normal content renders at 1.0; only genuinely-oversized text
            # shrinks — and it is ALWAYS drawn (tiny beats a blank spot). A single
            # call, so there is never a double-drawn overlay.
            floor = min(1.0, MIN_FONT / size)          # the "quality" 8pt floor
            pre_xrefs = set(page.get_contents())
            spare, scale = page.insert_htmlbox(rect, html, scale_low=0.05)
            _actualtext_wrap(doc, page, pre_xrefs, text)
            if scale >= 0.999:
                report['placed'] += 1
            elif scale >= floor - 0.001:
                report['shrunk'] += 1                  # readable shrink
            else:
                report['failed'] += 1                  # drawn, but below the 8pt floor
    if only_idx is not None:
        doc.select([only_idx])
    out = io.BytesIO(doc.tobytes())
    resp = send_file(out, mimetype='application/pdf', as_attachment=True,
                     download_name='translated.pdf')
    resp.headers['X-Build-Report'] = json.dumps(report)
    return resp


# ── DOCX lane ───────────────────────────────────────────────────────────────

def _eff(run, attr):
    v = getattr(run.font, attr)
    return bool(v) if v is not None else False


def docx_pars(d):
    def walk_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
                    yield from walk_tables(cell.tables)
    for p in d.paragraphs:
        yield p
    yield from walk_tables(d.tables)
    for s in d.sections:
        for part in (s.header, s.footer):
            for p in part.paragraphs:
                yield p
            yield from walk_tables(part.tables)


def docx_segments(par):
    segs = []
    for r in par.runs:
        if r._element.findall(qn('w:drawing')):
            continue
        if not r.text:
            continue
        key = (_eff(r, 'bold') or par.style.name.startswith('Heading'), _eff(r, 'italic'))
        if segs and segs[-1][1] == key:
            segs[-1][0] += r.text
        else:
            segs.append([r.text, key, r])
    return segs


def docx_marked(par):
    segs = docx_segments(par)
    if not segs or not ''.join(s[0] for s in segs).strip():
        return None, None
    base = max(segs, key=lambda s: len(s[0]))
    out = []
    for s in segs:
        if s[1] == base[1]:
            out.append(s[0])
        else:
            m = MARK.get(s[1])
            out.append('[[%s]]%s[[/%s]]' % (m, s[0], m) if m else s[0])
    return ''.join(out), base


@app.post('/v1/docx/extract')
def docx_extract():
    _auth()
    d = docx.Document(io.BytesIO(_file_bytes()))
    blocks = []
    for i, p in enumerate(docx_pars(d)):
        txt, _ = docx_marked(p)
        if txt is not None:
            blocks.append({'id': 'p%d' % i, 'text': txt})
    return jsonify({'engine': 'docx', 'blocks': blocks})


@app.post('/v1/docx/build')
def docx_build():
    _auth()
    tr = json.loads(request.form.get('translations') or '{}')
    if not isinstance(tr, dict) or not tr:
        abort(400, 'missing translations')
    d = docx.Document(io.BytesIO(_file_bytes()))
    pars = list(docx_pars(d))
    report = {'pars': 0, 'marker_fail': 0}
    for i, p in enumerate(pars):
        txt, base = docx_marked(p)
        if txt is None:
            continue
        pid = 'p%d' % i
        new = tr.get(pid)
        if not isinstance(new, str) or not new.strip():
            abort(422, 'missing id: ' + pid)
        for m in ('B', 'I', 'BI'):
            if txt.count('[[%s]]' % m) != new.count('[[%s]]' % m) or \
               txt.count('[[/%s]]' % m) != new.count('[[/%s]]' % m):
                report['marker_fail'] += 1
                new = M_RE.sub('', new)  # fail-safe: base format, never lose text
                break
        bfont = base[2].font
        bb, bi = base[1]
        for r in list(p.runs):
            if r._element.findall(qn('w:drawing')):
                continue
            r._element.getparent().remove(r._element)

        def add_run(seg_text, flags):
            add = p.add_run(seg_text)
            add.bold, add.italic = flags
            add.font.size = bfont.size
            add.font.name = bfont.name
            if bfont.color and bfont.color.rgb is not None:
                add.font.color.rgb = bfont.color.rgb

        pos = 0
        stack = []
        for m in M_RE.finditer(new):
            if m.start() > pos:
                add_run(new[pos:m.start()], stack[-1] if stack else (bb, bi))
            pos = m.end()
            if m.group(1):
                if stack:
                    stack.pop()
            else:
                k = m.group(2)
                stack.append((k in ('B', 'BI'), k in ('I', 'BI')))
        if pos < len(new):
            add_run(new[pos:], (bb, bi))
        report['pars'] += 1
    out = io.BytesIO()
    d.save(out)
    out.seek(0)
    resp = send_file(
        out, as_attachment=True, download_name='translated.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    resp.headers['X-Build-Report'] = json.dumps(report)
    return resp


@app.get('/healthz')
def healthz():
    return jsonify({'ok': True, 'pymupdf': fitz.__doc__.split(':')[0].strip()})


if __name__ == '__main__':
    app.run(host='127.0.0.1', port=8093)
